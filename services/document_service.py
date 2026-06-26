"""
Document Analyzer — PDF, DOCX, Excel, Email files
Ekstrak teks + URL dari dokumen, lalu analisis dengan NLP + URL checker
"""
import io
import re
import hashlib
from typing import Tuple, List
from loguru import logger


def extract_urls_from_text(text: str) -> List[str]:
    pattern = re.compile(
        r"http[s]?://(?:[a-zA-Z0-9$\-_.+!*'(),;:@&=/?#%])+"
    )
    return list(set(pattern.findall(text)))


def extract_text_pdf(file_bytes: bytes) -> Tuple[str, List[str]]:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
        full_text = "\n".join(text_parts)
        urls = extract_urls_from_text(full_text)
        return full_text[:5000], urls
    except Exception as e:
        logger.warning(f"PDF extract error: {e}")
        return "", []


def extract_text_docx(file_bytes: bytes) -> Tuple[str, List[str]]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        # Ambil juga teks dari tabel
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        full_text = "\n".join(paragraphs)
        urls = extract_urls_from_text(full_text)
        return full_text[:5000], urls
    except Exception as e:
        logger.warning(f"DOCX extract error: {e}")
        return "", []


def extract_text_txt(file_bytes: bytes) -> Tuple[str, List[str]]:
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        urls = extract_urls_from_text(text)
        return text[:5000], urls
    except Exception as e:
        logger.warning(f"TXT extract error: {e}")
        return "", []


def analyze_document(filename: str, file_bytes: bytes) -> dict:
    """
    Entry point untuk analisis dokumen.
    Return dict dengan extracted_text, urls, file_hash, file_type.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    file_hash = hashlib.sha256(file_bytes).hexdigest()

    if ext == "pdf":
        text, urls = extract_text_pdf(file_bytes)
        file_type = "PDF"
    elif ext in ("doc", "docx"):
        text, urls = extract_text_docx(file_bytes)
        file_type = "Word Document"
    elif ext in ("txt", "eml", "msg"):
        text, urls = extract_text_txt(file_bytes)
        file_type = "Text/Email"
    else:
        text, urls = "", []
        file_type = ext.upper() or "Unknown"

    return {
        "file_type": file_type,
        "file_hash": file_hash,
        "extracted_text": text,
        "urls_found": urls,
        "char_count": len(text),
        "url_count": len(urls),
    }
